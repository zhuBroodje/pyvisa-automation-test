#!/usr/bin/env python
# coding: utf-8

# In[2]:


import time
import matplotlib.pyplot as plt
import numpy as np
import yaml

from Oscilloscope import *
from PSU_6705 import *
from ElectronicLoad import *

from docx import Document
from docx.shared import Inches
from datetime import datetime
import os

class PowerTester:
    def __init__(self, config_file):
        self.oscilloscope = 0
        self.load = 0
        self.power_supply = 0
        self.config=0
        self.doc=Document()
        self.folder_path=0
        self.doc_path=0

        self.generate_record_file(config_file) 
        #Configure connection
        self.instrument_connection()
        self.instrument_init_config()


    def test_flow(self):
        #TODO
        self.load.on()
        self.power_supply.on()
        self.set_dc_test_mode()
        self.efficiency_current()
        self.load_regulation()
        #TODO
        #FIXME: channel selection
        self.set_ac_test_mode()
        self.swtich_frequency()
        self.ripple()
        self.power_sequence()


    #TODO
    def set_ac_test_mode(self):
        #FIXME: use filter?   
        #FIXME: selecte channel
        self.oscilloscope.set_coupling(2,"AC")
        #FIXME
        #self.oscilloscope.set_acquire_mode('SAMple')
      

    def swtich_frequency(self):
        self.oscilloscope.set_acquire_mode("SAMple")
        #FIXME channel
        self.oscilloscope.set_bandwidth(2,'TWEnty')
        #FIXME channel
        self.oscilloscope.set_trigger_a_edge_source(2)
        self.oscilloscope.set_trigger_a_edge_coupling('HFRej')
      
        #FIXME ONLY VALID FOR FIXED FREQUENCY
        self.doc.add_heading(f'Switch Frequency vs load current',level=2)
        sample_points,scale=self.generate_sample_points('switch frequency')
        sample_points=np.round(sample_points,self.load.get_precision())
        fo_list,co_list=[],[]
        self.load.on()
        for i in range(len(sample_points)):
            self.load.set_mode('c',sample_points[i])
            time.sleep(2)
            #FIXME how to scale
            self.oscilloscope.auto_scale(2,False)
            #FIXME how to get value
            #output_f=oscilloscope.get_measurement(5)
            #FIXME channel?
            output_f=self.oscilloscope.get_frequency(2)
            output_c=self.load.get_current_current()
            #Debug
            print(f"output {output_f}Hz {output_c}A")
            fo_list.append(float(output_f))
            co_list.append(float(output_c))

        plt.figure(figsize=(10, 5))
        plt.plot(co_list, fo_list)
        plt.grid(True)
        #FIXME scale?
        plt.xscale(scale)
        plt.yscale('log')
        plt.yticks(fo_list,fo_list)
        plt.xscale(scale)
        x_tick_values = np.linspace(0, max(co_list), 5)
        plt.xticks(x_tick_values, ['{:.2f}'.format(value) for value in x_tick_values])
        plt.title(f"Switch frequency vs Load current (Vin={self.config['DUT']['input']}V,Vout={self.config['DUT']['output']}V)")
        plt.ylabel('Freuquency(Hz)')
        plt.xlabel('I_LOAD (A)')
        plt.savefig('Switch frequency vs Load current.png')
        self.doc.add_picture('Switch frequency vs Load current.png', width=Inches(4))
        self.doc.save(self.doc_path)
            

    def ripple(self):
        #FIXME 
        self.oscilloscope.set_acquire_mode('SAMple')
        self.oscilloscope.set_trigger_a_edge_coupling('AC')
        self.oscilloscope.set_trigger_a_edge_source(2)
        self.oscilloscope.set_bandwidth(2,'TWEnty')
        self.oscilloscope.set_trigger_a_edge_coupling('HFRej')
        sample_points,scale=self.generate_sample_points('ripple')
        #FIXME how to get precision
        sample_points=np.round(sample_points,self.load.get_precision())
        vo_list,co_list,p_list=[],[],[]
        for i in range (len(sample_points)):
            self.load.set_mode('C',sample_points[i])
            time.sleep(2)
            #FIXME aaaaa
            self.oscilloscope.auto_scale(2,False)
            output_c=self.load.get_current_current()
            #FIXME how to get? wrong ripple measuremen!
            output_v=self.oscilloscope.measure_ripple(2)
            print(f"output {output_v} {output_c}")
            vo_list.append(output_v)
            co_list.append(output_c)
            #FIXME
            w,t=self.oscilloscope.get_waveform_data(2)
            data = np.vstack((w, t)).T
            name=f"waveform-current {output_c}A.txt"
            np.savetxt(name, data)
            #TODO show figures我好懒
            #md 烦了
        vo=self.config['DUT']['output']
        for v in vo_list:
            p=v/vo*100
            p_list.append(p)
        #p_list=vo_list/vo*100 #ripple/output DC value
        #debug
        print(vo_list,co_list,p_list)

    def power_sequence(self):
        #POWER ON
        #FIXME channel?
        self.oscilloscope.channel_off(3)
        self.oscilloscope.set_offset(1,0)
        self.oscilloscope.write('CH1:POSITION 0') 

        self.oscilloscope.write("TRIGGER:A:EDGE:SOURCE CH1")
        self.oscilloscope.write("TRIGGER:A:EDGE:SLOPE RISE")

        self.oscilloscope.set_offset(2,0)
        self.oscilloscope.set_t_scale(0.02)
        #self.oscilloscope.auto_y_scale(1)
        #self.oscilloscope.auto_y_scale(2)
        self.psu.off()
        self.load.off()

        #POWER OFF
        pass

    def generate_record_file(self,config_file):
        with open(config_file, 'r') as file:
            config = yaml.safe_load(file)
            self.config = config
        test_name=self.config['test_configuration']['test_name']
        current_date=datetime.now()
        test_date=current_date.strftime("%Y-%m-%d-%H%M")
        dut=self.config['DUT']['device_name']
        self.folder_path = f'{dut} {test_name} {test_date}'
        os.makedirs(self.folder_path)  
        self.doc.add_heading(f'{test_name}',level=1)
        self.doc.add_paragraph(f'Time: {test_date}\nDUT: {dut}')
        self.doc_path=f'{self.folder_path}/{dut} {test_name} {test_date}.docx'
        self.doc.save(self.doc_path)

    def instrument_connection(self):
        self.doc.add_heading(f'Test Instrument',level=2)
        #Oscilloscope connection
        self.doc.add_heading(f'Oscilloscope',level=3)
        self.oscilloscope = Oscilloscope(self.config['test_devices']['oscilloscope']['path'])
        self.doc.add_paragraph(f'idn:{self.oscilloscope.get_IDN()}')
        #Power Supply connection
        self.doc.add_heading(f'DC Power Supply',level=3)
        self.power_supply = PSU_6705(self.config['test_devices']['power_supply']['path'])
        self.doc.add_paragraph(f'idn:{self.power_supply.get_IDN()}')
        #Load connection
        self.doc.add_heading(f'Load',level=3)
        self.load = ElectronicLoad(self.config['test_devices']['load']['path'])
        self.doc.add_paragraph(f'idn:{self.load.get_IDN()}')
        #Save test intrument idn to file
        self.doc.save(self.doc_path)
       
    def instrument_init_config(self):
        self.load_init_config()
        self.osc_init_config()
        self.dcps_init_config()

    def load_init_config(self):
        #TODO should be configured by what?
        #load_settings = self.config['test_devices']['load']['settings']
        self.load.off()
        self.load.set_function('c')
        
    def osc_init_config(self):
        #initializing channels
        channel_settings = self.config['test_devices']['oscilloscope']['channel_settings']
        for channel, settings in channel_settings.items():
            self.oscilloscope.channel_on(channel)
            if 'coupling' in settings:
                self.oscilloscope.set_coupling(channel, settings['coupling'])
            else:
                self.oscilloscope.set_coupling(channel, 'DC')
            if 'bandwidth' in settings:
                self.oscilloscope.set_bandwidth(channel,settings['bandwidth'])
            else:
                self.oscilloscope.set_bandwidth(channel,'FULL')
        # adding measurements
        measurement_settings = self.config['test_devices']['oscilloscope']['measurement_settings']
        for measurement, settings in measurement_settings.items():
            self.oscilloscope.add_measurement(measurement,settings['channel'], settings['type'])
       
    def dcps_init_config(self):
        settings=self.config['test_devices']['power_supply']['settings']
        self.power_supply.off()
        self.power_supply.set_value(1,'v',settings['input_voltage'])
        self.power_supply.set_value(1,'c',settings['input_current'])
        self.power_supply.set_value(2,'v',settings['testbord_supply'])
        self.power_supply.set_value(2,'c',settings['testbord_supply_current'])

    def set_dc_test_mode(self):
        #FIXME select channel
        self.oscilloscope.set_t_scale(0.002)
        self.oscilloscope.auto_y_scale(1)
        self.oscilloscope.auto_y_scale(2)
        #no multimeter is used, so use oscilloscope and electronicload to read value for now
        pass # already initialized by  instrument_init_config

    def end_test(self):
        self.load.load.close()
        self.oscilloscope.scope.close()
        self.power_supply.close()
        print("Test ended")
        
    
    def efficiency_current(self):
        self.doc.add_heading(f'Effiency vs load current',level=2)
        #Note : for debug
        print("efficiency vs current")
        sample_points,scale= self.generate_sample_points('efficiency')
        sample_points=np.round(sample_points,self.load.get_precision()) # Notes: should get precision from instruments, 4 is only for this load
        
        vo_list, co_list, vi_list, ci_list = [], [], [], []
        #Note: Maybe logging?
        for i in range (len(sample_points)):
            self.load.set_mode('C',sample_points[i])
            time.sleep(2) #wait for value stablized
            output_v=self.oscilloscope.get_measurement(2)
            #output_v=self.load.get_current_voltage()
            output_c=self.load.get_current_current()
            #Note : for debug
            print(f"output {output_v} {output_c}")
            vo_list.append(output_v)
            co_list.append(output_c)
            
            input_status=self.power_supply.status_inquire()
            #input_v=input_status.get('ch1_out_voltage')
            # TODO: how to scale?
            # FIXME: how to get the value?
            input_v=self.oscilloscope.get_measurement(1)
            input_c=input_status.get('ch1_out_current')
            # Note: for debug
            print(f"input {input_v} {input_c}")
            vi_list.append(input_v)
            ci_list.append(input_c)
        #FIXME 
        #incase input_c read 0, truncate those samples
        vo_array, co_array, vi_array, ci_array = np.array(vo_list), np.array(co_list), np.array(vi_list), np.array(ci_list)
        po_array, pi_array = vo_array * co_array, vi_array * ci_array
        efficiency=po_array/pi_array *100
         # Note: for debug
        print("po",po_array)
        print("pi",pi_array)
        print("eff",efficiency)

        plt.figure(figsize=(10, 5))
        plt.plot(co_list, efficiency)
        plt.xscale(scale)     
        plt.grid(True)
        plt.title(f"Efficiency vs Load current (Vin={self.config['DUT']['input']}V,Vout={self.config['DUT']['output']}V)")
        plt.ylabel('Efficiency(%)')
        plt.xlabel('I_LOAD (A)')
        plt.savefig('Efficiency vs Load current.png')
        self.doc.add_picture('Efficiency vs Load current.png', width=Inches(4))
        self.doc.save(self.doc_path)

    def generate_points(self,start,end,num,scale):
        if scale == 'log':
            return np.logspace(np.log10(start), np.log10(end), num)  
        elif scale=='linear' :
            return np.linspace(start, end, num)
        else:
            return np.linspace(start, end, num)
    def generate_sample_points(self,update_config=''):
        test_configuration = self.config['test_configuration'].copy()
        settings=test_configuration['default']
        if update_config in test_configuration and test_configuration[update_config]:
            for key, value in test_configuration[update_config].items():
                if key is not None:
                    settings[key]=value
        min_load, max_load, sample_num,scale = [settings.get(key) for key in ['min_load', 'max_load', 'sample_num','sample_scale']]
        print(f"min_load{min_load}, max_load{max_load}, sample_num{sample_num},scale{scale}")
        sample_points=self.generate_points(min_load,max_load,sample_num,scale)
        return sample_points,scale    
        
    def load_regulation(self):
        self.doc.add_heading(f'Load regulation',level=2)
        v_list,c_list=[],[]
        sample_points,scale=self.generate_sample_points('load_regulation')
        sample_points=np.round(sample_points,self.load.get_precision())
        for i in range (len(sample_points)):
            self.load.set_mode('C', sample_points[i] )
            time.sleep(2)
            v=self.load.get_current_voltage()
            c=self.load.get_current_current()
            v_list.append(v)
            c_list.append(c)
            print(f"Voltage {v} Current {c}")
        #Debug
        print(v_list,c_list)
        plt.figure(figsize=(10, 5))
        plt.plot(c_list, v_list)
        plt.grid(True)
        plt.xscale(scale)
        title=f"Load regulation(Vin={self.config['DUT']['input']}V,Vout={self.config['DUT']['output']}V)"
        plt.title(title)
        plt.ylabel('Vout (V)')
        plt.xlabel('I_LOAD (A)')
        plt.savefig('Load regulation.png')
        self.doc.add_picture('Load regulation.png', width=Inches(4))
        self.doc.save(self.doc_path)
       



# In[ ]:
